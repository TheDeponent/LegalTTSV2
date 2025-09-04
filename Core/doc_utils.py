# doc_utils.py - Unified Document Utilities for LegalTTSV2
# Provides all document-related utilities: DOCX extraction, chunking, file normalization,
# RTF/PDF conversion, and preprocessing. All logic is GUI-agnostic and logging is consistent.

import os
import re
import pypandoc
import shutil
import urllib.parse
import pythoncom
import win32com.client
from docx import Document
from Core.constants import MAX_CHUNK_LENGTH
from Core.llm_handler import log



def extract_text_from_docx(docx_path):
    # Extract all text from a .docx file and return as a single string.
    try:
        doc = Document(docx_path)
        return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        log(f"Error extracting text from docx file: {e}")
        return ""

def extract_paragraph_chunks(docx_path):
    # Extract non-empty paragraphs from a .docx file as a list of strings (chunks).
    doc = Document(docx_path)
    chunks = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            chunks.append(text)
    return chunks

def split_long_paragraphs(paragraphs, max_length=MAX_CHUNK_LENGTH):
    # Split paragraphs longer than max_length into smaller chunks, breaking at sentence boundaries if possible.
    chunks = []
    for para in paragraphs:
        if len(para) <= max_length:
            chunks.append(para)
        else:
            log(f"Splitting long paragraph of length {len(para)}")
            start = 0
            while start < len(para):
                if len(para) - start <= max_length:
                    chunks.append(para[start:].strip())
                    break
                split_idx = start + max_length
                match = re.search(r'[.!]', para[split_idx:])
                if match:
                    end = split_idx + match.end()
                    chunks.append(para[start:end].strip())
                    start = end
                else:
                    chunks.append(para[start:start+max_length].strip())
                    start += max_length
    return chunks

def normalize_path(path):
    # Normalize a file path by removing URL encoding and normalizing slashes.
    return os.path.abspath(os.path.normpath(urllib.parse.unquote(path)))


def convert_rtf_to_docx(rtf_path):
    # Convert an RTF file to DOCX using pypandoc. Returns the path to the new DOCX file.
    if not pypandoc:
        log("pypandoc is required for RTF to DOCX conversion.'.")
        raise ImportError("pypandoc is required for RTF to DOCX conversion.'.")
    docx_path = os.path.splitext(rtf_path)[0] + ".docx"
    pypandoc.convert_file(rtf_path, 'docx', outputfile=docx_path)
    if not os.path.exists(docx_path):
        log(f"Failed to convert {rtf_path} to DOCX.")
        raise RuntimeError(f"Failed to convert {rtf_path} to DOCX.")
    log(f"Converted RTF to DOCX: {docx_path}")
    return docx_path


def convert_to_docx(input_path):
    # Convert an input file (RTF, PDF, or DOCX) to DOCX if needed.
    ext = os.path.splitext(input_path)[1].strip().lower()
    if ext == ".docx":
        log(f"Input is already DOCX: {input_path}")
        return input_path
    elif ext == ".rtf":
        return convert_rtf_to_docx(input_path)
    elif ext == ".pdf":
        return process_pdf(input_path)
    elif ext == ".txt":
        # Convert .txt to .docx by reading text and writing to a new docx file
        from docx import Document
        doc = Document()
        with open(input_path, 'r', encoding='utf-8') as f:
            for line in f:
                doc.add_paragraph(line.rstrip())
        docx_path = os.path.splitext(input_path)[0] + ".docx"
        doc.save(docx_path)
        log(f"Converted TXT to DOCX: {docx_path}")
        return docx_path
    else:
        log(f"Unsupported file type for conversion: {input_path} (extension: '{ext}')")
        raise ValueError(f"Unsupported file type for conversion: {input_path} (extension: '{ext}')")


def pdf_to_docx(pdf_path: str, docx_path: str) -> None:
    # Convert a PDF file to a Word (.docx) file using Microsoft Word automation (if available),
    # otherwise falls back to pdf2docx library. Logs all major events.
    try:
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        pdf_path = os.path.normpath(os.path.abspath(pdf_path))
        docx_path = os.path.normpath(os.path.abspath(docx_path))
        try:
            doc = word.Documents.Open(pdf_path)
            doc.SaveAs(docx_path, FileFormat=16)  # 16 = wdFormatDocumentDefault (docx)
            doc.Close()
        finally:
            word.Quit()
        log(f"Converted PDF to DOCX using Word: {docx_path}")
    except Exception as e:
        log(f"Word automation failed: {e}. Falling back to pdf2docx for conversion.")
        from pdf2docx import Converter
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        log(f"Converted PDF to DOCX using pdf2docx: {docx_path}")


def preprocess_docx(docx_path: str) -> None:
    # Preprocess the docx file in-place: removes images, numbers, dashes, phone numbers, emails, and paragraph numbers.
    doc = Document(docx_path)
    for shape in doc.inline_shapes:
        shape._inline.getparent().remove(shape._inline)
    for p in doc.paragraphs:
        for run in p.runs:
            if hasattr(run, 'element') and hasattr(run.element, 'xpath'):
                for drawing in run.element.xpath('.//w:drawing'):
                    parent = drawing.getparent()
                    try:
                        if parent is not None:
                            parent.remove(drawing)
                    except ValueError:
                        log("Warning: drawing element is not a child of its parent, skipping removal.")
    num_in_brackets = re.compile(r'\[\s*\d+\s*\]')
    hanging_dash = re.compile(r' - ')
    phone = re.compile(r'\b\+?\d[\d\s\-]{7,}\d\b')
    email = re.compile(r'[\w\.-]+@[\w\.-]+')
    para_number = re.compile(r'^(\(?[a-zA-Z0-9]+\)?[\.|\)]\s*)')
    for p in doc.paragraphs:
        text = p.text
        text = num_in_brackets.sub('', text)
        text = hanging_dash.sub(' ', text)
        text = phone.sub('', text)
        text = email.sub('', text)
        text = para_number.sub('', text)
        p.text = text.strip()
    doc.save(docx_path)
    log(f"Preprocessed DOCX: {docx_path}")


def process_pdf(pdf_path: str, temp_dir=None) -> str:
    # Convert a PDF to DOCX, save a preprocessed copy, preprocess the DOCX, and return the path to the processed file.
    if temp_dir is None:
        temp_dir = os.path.dirname(pdf_path)
    base = os.path.splitext(os.path.basename(pdf_path))[0]
    docx_path = os.path.join(temp_dir, f"{base}_processed.docx")
    pdf_to_docx(pdf_path, docx_path)
    preproc_copy = os.path.join(temp_dir, f"{base}_original.docx")
    shutil.copyfile(docx_path, preproc_copy)
    preprocess_docx(docx_path)
    log(f"Processed PDF to DOCX: {docx_path}")
    return docx_path
def docx_to_txt(docx_path, txt_path=None):
    # Convert a DOCX file to plain TXT (one paragraph per line).
    doc = Document(docx_path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    txt_content = '\n'.join(lines)
    if txt_path:
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(txt_content)
        log(f"Converted DOCX to TXT: {txt_path}")
        return txt_path
    return txt_content

def rtf_to_txt(rtf_path, txt_path=None):
    # Convert RTF to DOCX, then to TXT.
    docx_path = convert_rtf_to_docx(rtf_path)
    return docx_to_txt(docx_path, txt_path)

def ensure_ollama_text(input_path):
    # For Ollama: if TXT, read and return text. If PDF/RTF, convert to DOCX, extract text. If DOCX, extract text.
    ext = os.path.splitext(input_path)[1].strip().lower()
    if ext == '.txt':
        with open(input_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif ext == '.docx':
        return extract_text_from_docx(input_path)
    elif ext == '.rtf' or ext == '.pdf':
        docx_path = convert_to_docx(input_path)
        return extract_text_from_docx(docx_path)
    else:
        raise ValueError(f"Unsupported file type for Ollama: {input_path} (extension: '{ext}')")