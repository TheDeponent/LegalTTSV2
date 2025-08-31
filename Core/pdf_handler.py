
# ============================================================================
# pdf_handler.py - PDF to DOCX Conversion and Preprocessing for LegalTTSV2
#
# This module provides functions for converting PDF files to DOCX using Microsoft Word
# automation, and for preprocessing DOCX files to remove unwanted content (images, numbers,
# emails, phone numbers, etc). It is used by the main workflow to ensure documents are in
# a clean, processable format for LLM and TTS operations.
# ============================================================================

import os
import re
from typing import Optional
from docx import Document
import win32com.client
import shutil

def pdf_to_docx(pdf_path: str, docx_path: str) -> None:
    # Converts a PDF file to a Word (.docx) file using Microsoft Word automation (win32com.client).
    # Used by process_pdf to prepare documents for further processing.
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    # Normalize paths for Word
    pdf_path = os.path.normpath(os.path.abspath(pdf_path))
    docx_path = os.path.normpath(os.path.abspath(docx_path))
    try:
        doc = word.Documents.Open(pdf_path)
        doc.SaveAs(docx_path, FileFormat=16)  # 16 = wdFormatDocumentDefault (docx)
        doc.Close()
    finally:
        word.Quit()

def preprocess_docx(docx_path: str) -> None:
    # Preprocesses the docx file in-place: removes images, paragraph numbering/lettering, [numbers],
    # hanging dashes, phone numbers, and emails. Used by process_pdf to clean up documents for LLM/TTS.
    doc = Document(docx_path)
    # Remove all images (inline shapes, pictures, drawings)
    for shape in doc.inline_shapes:
        # Remove the shape by clearing the paragraph text
        shape._inline.getparent().remove(shape._inline)
    for p in doc.paragraphs:
        # Remove images in runs (w:drawing)
        for run in p.runs:
            if hasattr(run, 'element') and hasattr(run.element, 'xpath'):
                for drawing in run.element.xpath('.//w:drawing'):
                    parent = drawing.getparent()
                    try:
                        if parent is not None:
                            parent.remove(drawing)
                    except ValueError:
                        print("Warning: drawing element is not a child of its parent, skipping removal.")
    # Preprocessing regexes for numbers, dashes, phone numbers, emails, and paragraph numbers
    num_in_brackets = re.compile(r'\[\s*\d+\s*\]')
    hanging_dash = re.compile(r' - ')
    phone = re.compile(r'\b\+?\d[\d\s\-]{7,}\d\b')
    email = re.compile(r'[\w\.-]+@[\w\.-]+')
    para_number = re.compile(r'^(\(?[a-zA-Z0-9]+\)?[\.|\)]\s*)')
    for p in doc.paragraphs:
        text = p.text
        text = num_in_brackets.sub('', text)
        text = hanging_dash.sub(' ', text)
        text = phone.sub('', text)
        text = email.sub('', text)
        text = para_number.sub('', text)
        p.text = text.strip()
    doc.save(docx_path)

def process_pdf(pdf_path: str, temp_dir: Optional[str] = None) -> str:
    # Converts a PDF to DOCX, saves a preprocessed copy, preprocesses the DOCX, and returns the path to the processed file.
    # pdf_to_docx is used to convert the PDF; preprocess_docx is used to clean up the DOCX.
    if temp_dir is None:
        temp_dir = os.path.dirname(pdf_path)
    base = os.path.splitext(os.path.basename(pdf_path))[0]
    docx_path = os.path.join(temp_dir, f"{base}_processed.docx")
    pdf_to_docx(pdf_path, docx_path)
    # Save a copy before preprocessing for reference
    preproc_copy = os.path.join(temp_dir, f"{base}_original.docx")
    shutil.copyfile(docx_path, preproc_copy)
    preprocess_docx(docx_path)
    return docx_path
