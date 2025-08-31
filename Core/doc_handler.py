# ============================================================================
# doc_handler.py - DOCX Document Handler for LegalTTSV2
#
# This module provides functions for extracting text and paragraphs from .docx files.
# It is used by the main application workflow and other modules to load document
# content for LLM and TTS processing. All DOCX reading and parsing logic is handled here.
# ============================================================================

from docx import Document
from Core.constants import MAX_CHUNK_LENGTH

def extract_text_from_docx(docx_path):
    # Extracts all text from a .docx file and returns it as a single string.
    # docx.Document is used to load and parse the DOCX file.
    try:
        doc = Document(docx_path)
        return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        print(f"Error extracting text from docx file: {e}")
        return ""

def extract_paragraph_chunks(docx_path):
    # Extracts non-empty paragraphs from a .docx file as a list of strings (chunks).
    # docx.Document is used to load and parse the DOCX file.
    doc = Document(docx_path)
    chunks = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            chunks.append(text)
    return chunks
